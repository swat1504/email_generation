import streamlit as st
import os
import re
import base64
import zipfile
import tempfile
import shutil
import pandas as pd
import json
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

SCOPES = ['https://www.googleapis.com/auth/gmail.compose']


# ---------------- Utilities ----------------

def normalize_branch_name(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9]', '', name).upper()


# ---------------- Dynamic Signature ----------------

def get_signature_html(sender_email):

    if sender_email == "ambikanathp@gmail.com":
        name = "Ambikanath Parida"

    elif sender_email == "janakinathp@gmail.com":
        name = "Janakinath Parida"

    elif sender_email == "swatiparida1504@gmail.com":
        name = "Swati Parida"

    else:
        name = ""

    return f"""
    <br><br>
    <p>With Warm Regards,</p>
    <p style="font-size:20px; font-weight:bold; color:#5aa05a;">
        {name}
    </p>
    <p style="font-size:18px; font-weight:bold; color:#2f3df4;">
        Gayattree Docutechs Private Limited
    </p>
    """


# ---------------- Gmail Service ----------------

def get_gmail_service(sender_email):

    if sender_email == "ambikanathp@gmail.com":
        token_json = st.secrets["token_ambikanath"]

    elif sender_email == "janakinathp@gmail.com":
        token_json = st.secrets["token_janakinath"]

    elif sender_email == "swatiparida1504@gmail.com":
        token_json = st.secrets["token_swati"]

    else:
        raise Exception("Unknown sender email")

    creds = Credentials.from_authorized_user_info(
        json.loads(token_json),
        SCOPES
    )

    return build("gmail", "v1", credentials=creds)


# ---------------- Word → HTML Conversion ----------------

def extract_subject_and_body(docx_path):

    doc = Document(docx_path)
    subject = ""
    html_parts = []

    # ---- Extract subject ----
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith("subject"):
            subject = p.text.replace("Subject:", "").strip()
            break

    # ---- Convert paragraphs with proper line breaks ----
    for p in doc.paragraphs:

        text = p.text.strip()

        if not text:
            html_parts.append("<br>")
            continue

        # Convert internal Word line breaks to HTML
        text = text.replace("\n", "<br>")

        html_parts.append(
            f"<p style='margin:0 0 10px 0; line-height:1.6;'>{text}</p>"
        )

    # ---- Convert tables ----
    for table in doc.tables:

        html_parts.append("<br>")
        html_parts.append(
            "<table border='1' cellpadding='6' cellspacing='0' "
            "style='border-collapse: collapse; width:100%;'>"
        )

        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_text = cell.text.replace("\n", "<br>")
                html_parts.append(
                    f"<td style='text-align:left;'>{cell_text}</td>"
                )
            html_parts.append("</tr>")

        html_parts.append("</table>")
        html_parts.append("<br>")

    return subject, "".join(html_parts)


# ---------------- Create Draft ----------------

def create_draft_email(service, sender, to, cc, bcc, subject, body_html, attachments):

    message = MIMEMultipart("alternative")
    message["to"] = to
    message["from"] = sender
    message["subject"] = subject

    if cc:
        message["cc"] = cc
    if bcc:
        message["bcc"] = bcc

    message.attach(MIMEText(body_html, "html"))

    for file_path in attachments:
        with open(file_path, "rb") as f:
            part = MIMEApplication(
                f.read(),
                Name=os.path.basename(file_path)
            )
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        message.attach(part)

    raw = base64.urlsafe_b64encode(message.as_bytes()).decode()

    service.users().drafts().create(
        userId="me",
        body={"message": {"raw": raw}}
    ).execute()


# ---------------- Core Processing ----------------

def process_emails(folder_path, sender_email, cc_email, bcc_email):

    logs = []
    service = get_gmail_service(sender_email)

    # ---- Find Excel ----
    excel_files = []

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith((".xls", ".xlsx")):
                excel_files.append(os.path.join(root, file))

    if not excel_files:
        return ["❌ No Excel file found inside uploaded ZIP."]

    excel_path = excel_files[0]
    logs.append(f"📄 Excel Found: {os.path.basename(excel_path)}")

    # ---- Detect Correct Sheet ----
    xls = pd.ExcelFile(excel_path)

    selected_df = None

    for sheet in xls.sheet_names:

        temp_df = pd.read_excel(xls, sheet_name=sheet)

        cols = [c.strip().upper() for c in temp_df.columns]

        # Normalize column names
        branch_col = None
        email_col = None

        for c in cols:
            if "BRANCH" in c:
                branch_col = c
            if "EMAIL" in c:
                email_col = c

        if branch_col and email_col:
            selected_df = temp_df
            break


    if selected_df is None:
        return ["❌ No sheet found containing BRANCH and EMAIL columns"]


    # Standardize column names internally
    selected_df.columns = [c.strip().upper() for c in selected_df.columns]

    branch_column = [c for c in selected_df.columns if "BRANCH" in c][0]
    email_column = [c for c in selected_df.columns if "EMAIL" in c][0]

    df = selected_df[[branch_column, email_column]].copy()

    df.columns = ["Branch Name", "Email"]

    df["Branch Key"] = df["Branch Name"].astype(str).apply(normalize_branch_name)

    # ---- Detect Required Folders ----
    letters_path = None
    invoices_path = None

    for root, dirs, files in os.walk(folder_path):
        for d in dirs:
            if d.upper() == "LETTERS":
                letters_path = os.path.join(root, d)
            if d.upper() == "BRANCHWISE INVOICES":
                invoices_path = os.path.join(root, d)

    if not letters_path:
        return ["❌ LETTERS folder not found."]
    if not invoices_path:
        return ["❌ BRANCHWISE INVOICES folder not found."]

    # ---- Process Each Letter ----
    for file in os.listdir(letters_path):

        if not file.endswith("_letter.docx"):
            continue

        branch_raw = file.replace("_letter.docx", "")
        branch_key = normalize_branch_name(branch_raw)

        row = df[df["Branch Key"] == branch_key]

        if row.empty:
            logs.append(f"⚠️ No email for {branch_raw}")
            continue

        recipient_email = row["Email"].values[0]

        invoice_folder = None

        for folder in os.listdir(invoices_path):
            if normalize_branch_name(folder) == branch_key:
                invoice_folder = os.path.join(invoices_path, folder)
                break

        if not invoice_folder:
            logs.append(f"⚠️ No invoices for {branch_raw}")
            continue

        attachments = [
            os.path.join(invoice_folder, f)
            for f in os.listdir(invoice_folder)
            if os.path.isfile(os.path.join(invoice_folder, f))
        ]

        subject, body_html = extract_subject_and_body(
            os.path.join(letters_path, file)
        )

        body_html += get_signature_html(sender_email)

        create_draft_email(
            service,
            sender_email,
            recipient_email,
            cc_email,
            bcc_email,
            subject,
            body_html,
            attachments
        )

        logs.append(f"✅ Draft created for {recipient_email}")

    return logs


# ---------------- Streamlit UI ----------------

st.set_page_config(page_title="Bulk Gmail Draft Generator", layout="centered")

st.title("📧 Bulk Gmail Draft Generator")

uploaded_zip = st.file_uploader(
    "📁 Upload your main folder as ZIP",
    type=["zip"]
)

st.markdown("---")

sender_email = st.selectbox(
    "📧 Select Sender Email",
    [
        "ambikanathp@gmail.com",
        "janakinathp@gmail.com",
        "swatiparida1504@gmail.com"
    ]
)

cc_email = st.text_input("📧 CC (optional)")
bcc_email = st.text_input("📧 BCC (optional)")

col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    generate_button = st.button(
        "🚀 Generate Draft Emails",
        use_container_width=True
    )

if generate_button:

    if not uploaded_zip:
        st.error("Please upload ZIP file.")
    else:

        with tempfile.TemporaryDirectory() as temp_dir:

            zip_path = os.path.join(temp_dir, "uploaded.zip")

            with open(zip_path, "wb") as f:
                f.write(uploaded_zip.read())

            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Remove macOS metadata
            for item in os.listdir(temp_dir):
                if item.startswith("__MACOSX"):
                    shutil.rmtree(os.path.join(temp_dir, item))

            items = [i for i in os.listdir(temp_dir) if i != "uploaded.zip"]

            if len(items) == 1 and os.path.isdir(os.path.join(temp_dir, items[0])):
                extracted_folder = os.path.join(temp_dir, items[0])
            else:
                extracted_folder = temp_dir

            with st.spinner("Generating drafts..."):
                logs = process_emails(
                    extracted_folder,
                    sender_email,
                    cc_email,
                    bcc_email
                )

            st.success("Process Completed!")

            for log in logs:
                st.write(log)
